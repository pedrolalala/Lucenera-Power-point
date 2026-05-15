"""
SharePoint Client para integração com Microsoft Graph API
Versão SIMPLIFICADA - Busca por código interno em pasta única
"""

import requests
import msal
import json
import os
import re
import sys
import logging
from typing import List, Dict, Optional, Tuple
import io
from PIL import Image
from dotenv import load_dotenv

# Configurar encoding UTF-8 para console Windows
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except:
        pass  # Se falhar, usar logging normal

# Carregar variáveis de ambiente do arquivo .env
load_dotenv()

# Configurar logger principal
def setup_sharepoint_logger():
    """Configura logger para SharePoint com formato personalizado"""
    logger = logging.getLogger('sharepoint_client')
    
    # Evitar duplicação de handlers
    if logger.handlers:
        return logger
    
    logger.setLevel(logging.INFO)
    
    # Handler para arquivo
    file_handler = logging.FileHandler('sharepoint_operations.log', encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    
    # Handler para console  
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    
    # Formato personalizado
    formatter = logging.Formatter(
        '%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)
    
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    return logger

# Inicializar logger
sharepoint_logger = setup_sharepoint_logger()

class SharePointClient:
    """Cliente SIMPLIFICADO para acessar arquivos do SharePoint via Microsoft Graph API"""
    
    # DRIVE ID ESPECÍFICO DO DRIVE "FICHAS TÉCNICAS"
    DRIVE_ID_FICHAS_TECNICAS = "b!hn714jDw0EaIQA68sMO3by9qF5mSQddKsXvCP4jIhNY4GELowLp8QpiQqx6Bdznf"
    
    # CAMINHO FIXO ÚNICO NO DRIVE DE FICHAS TÉCNICAS
    PASTA_FICHAS_UNICA = "Ficha técnica processada/Ficha Técnica Renomeado"
    
    def __init__(self):
        # Logger
        self.logger = sharepoint_logger
        
        # Credenciais Microsoft Graph (carregadas do .env)
        self.client_id = os.getenv('SHAREPOINT_CLIENT_ID')
        self.tenant_id = os.getenv('SHAREPOINT_TENANT_ID')
        self.client_secret = os.getenv('SHAREPOINT_CLIENT_SECRET')
        
        # Validar se credenciais foram carregadas
        if not all([self.client_id, self.tenant_id, self.client_secret]):
            missing = []
            if not self.client_id: missing.append('SHAREPOINT_CLIENT_ID')
            if not self.tenant_id: missing.append('SHAREPOINT_TENANT_ID') 
            if not self.client_secret: missing.append('SHAREPOINT_CLIENT_SECRET')
            
            raise EnvironmentError(
                f"Credenciais SharePoint não encontradas no arquivo .env: {', '.join(missing)}\n"
                f"Verifique se o arquivo .env existe e contém as variáveis necessárias."
            )
        
        # Scopes e endpoints
        self.authority = os.getenv('SHAREPOINT_AUTHORITY', f"https://login.microsoftonline.com/{self.tenant_id}")
        self.scope = [os.getenv('SHAREPOINT_SCOPE', "https://graph.microsoft.com/.default")]
        self.graph_url = os.getenv('SHAREPOINT_BASE_URL', "https://graph.microsoft.com/v1.0")
        
        # Cache de token
        self.access_token = None
        
        # Configurações SharePoint (carregadas do .env)
        self.site_name = os.getenv('SHAREPOINT_DRIVE_NAME', "LUCENERA PROJETOS")
        
        print(f"[OK] SharePointClient inicializado")
        print(f"[CAMINHO] Pasta fixa: {self.PASTA_FICHAS_UNICA}")
        
        # Log inicial de configuração
        self.logger.info(f"SharePoint Client inicializado com caminho fixo: {self.PASTA_FICHAS_UNICA}")
    
    def _get_access_token(self) -> str:
        """Obtém token de acesso para Microsoft Graph"""
        if self.access_token:
            return self.access_token
            
        try:
            # Criar aplicação MSAL
            app = msal.ConfidentialClientApplication(
                client_id=self.client_id,
                client_credential=self.client_secret,
                authority=self.authority
            )
            
            # Obter token
            result = app.acquire_token_for_client(scopes=self.scope)
            
            if result and "access_token" in result:
                self.access_token = result["access_token"]
                return self.access_token
            else:
                error_desc = result.get('error_description') if result else 'Resposta vazia da API'
                raise Exception(f"Erro na autenticação: {error_desc}")
                
        except Exception as e:
            raise Exception(f"Falha ao obter token: {str(e)}")
    
    def _make_graph_request(self, endpoint: str) -> Dict:
        """Faz requisição para Microsoft Graph API"""
        token = self._get_access_token()
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        url = f"{self.graph_url}{endpoint}"
        response = requests.get(url, headers=headers)
        
        if response.status_code == 200:
            return response.json()
        else:
            raise Exception(f"Erro na requisição: {response.status_code} - {response.text}")
    
    def get_site_id(self) -> str:
        """Obtém ID do site SharePoint"""
        try:
            # Buscar site por nome
            search_endpoint = f"/sites?search={self.site_name}"
            result = self._make_graph_request(search_endpoint)
            
            # Procurar site exato
            for site in result.get("value", []):
                if site.get("displayName") == self.site_name:
                    return site["id"]
                    
            raise Exception(f"Site '{self.site_name}' não encontrado")
            
        except Exception as e:
            raise Exception(f"Erro ao obter site ID: {str(e)}")
    
    def get_folder_id(self, site_id: str, folder_path: str) -> str:
        """Obtém ID de uma pasta no SharePoint"""
        try:
            # Buscar pasta por caminho
            folder_endpoint = f"/sites/{site_id}/drive/root:/{folder_path}"
            result = self._make_graph_request(folder_endpoint)
            
            return result["id"]
            
        except Exception as e:
            raise Exception(f"Erro ao obter pasta '{folder_path}': {str(e)}")
    
    
    def search_files_by_code(self, codigo_interno: str) -> List[Dict]:
        """
        Busca arquivos .docx no SharePoint por CÓDIGO INTERNO na pasta fixa
        
        Args:
            codigo_interno: Código interno do produto (ex: "10289", "10539")
            
        Returns:
            Lista de arquivos .docx encontrados com score de relevância
        """
        try:
            self.logger.info(f"🔍 Buscando arquivo para código interno: '{codigo_interno}'")
            
            # 1. Usar DRIVE ID específico (não precisa buscar site dinamicamente)
            self.logger.info(f"📂 Pasta fixa: {self.PASTA_FICHAS_UNICA}")
            
            # 2. Carregar arquivos da pasta usando DRIVE_ID
            search_endpoint = f"/drives/{self.DRIVE_ID_FICHAS_TECNICAS}/root:/{self.PASTA_FICHAS_UNICA}:/children"
            
            try:
                result = self._make_graph_request(search_endpoint)
                self.logger.info(f"[OK] Pasta SharePoint carregada: {len(result.get('value', []))} itens encontrados")
                
            except Exception as e:
                self.logger.error(f"[ERRO] Erro ao carregar pasta SharePoint: {str(e)}")
                return []
            
            # 3. Filtrar arquivos relevantes (.docx para fichas, .jpg/.png/.pdf para bulas)
            EXTENSOES_VALIDAS = ['.docx', '.jpg', '.jpeg', '.png', '.pdf']
            arquivos_relevantes = []
            for item in result.get("value", []):
                nome = item.get("name", "")
                nome_lower = nome.lower()
                if any(nome_lower.endswith(ext) for ext in EXTENSOES_VALIDAS):
                    arquivos_relevantes.append(item)
            
            self.logger.info(f"📄 {len(arquivos_relevantes)} arquivo(s) encontrados na pasta (docx/jpg/png/pdf)")
            
            # 4. Buscar match por código interno
            arquivos_encontrados = []
            for item in arquivos_relevantes:
                nome = item.get("name", "")
                nome_sem_ext = nome.rsplit('.', 1)[0]  # Remove extensão
                
                # Extrair código do início do nome (antes de underscore ou espaço)
                # Ex: "8223_BULA_01" → "8223", "9923" → "9923", "10289 - Descrição" → "10289"
                codigo_extraido = nome_sem_ext.split('_')[0].split(' ')[0].split('-')[0].strip()
                
                # Match por código extraído (case-insensitive)
                codigo_interno_upper = str(codigo_interno).upper()
                codigo_extraido_upper = codigo_extraido.upper()
                nome_sem_ext_upper = nome_sem_ext.upper()
                
                # Match EXATO - Código extraído corresponde ao código buscado
                if codigo_extraido_upper == codigo_interno_upper:
                    arquivo_info = {
                        'name': nome,
                        'id': item.get('id'),
                        'download_url': item.get('@microsoft.graph.downloadUrl'),
                        'web_url': item.get('webUrl'), 
                        'score': 100,  # Match perfeito
                        'match_method': 'codigo_extraido_exato',
                        'size': item.get('size', 0),
                        'last_modified': item.get('lastModifiedDateTime'),
                        'company_folder': self.PASTA_FICHAS_UNICA,
                        'type': self._detectar_tipo_arquivo(nome),
                        'is_bula': self._is_bula_file(nome)
                    }
                    
                    arquivos_encontrados.append(arquivo_info)
                    self.logger.info(f"🎯 Match EXATO encontrado: {nome} (código extraído: {codigo_extraido})")
                    
                # Match PARCIAL - Nome completo sem extensão corresponde (ex: "10289" == "10289")
                elif nome_sem_ext_upper == codigo_interno_upper:
                    arquivo_info = {
                        'name': nome,
                        'id': item.get('id'),
                        'download_url': item.get('@microsoft.graph.downloadUrl'),
                        'web_url': item.get('webUrl'),
                        'score': 95,  # Match nome completo
                        'match_method': 'nome_completo_exato',
                        'size': item.get('size', 0),
                        'last_modified': item.get('lastModifiedDateTime'),
                        'company_folder': self.PASTA_FICHAS_UNICA,
                        'type': self._detectar_tipo_arquivo(nome),
                        'is_bula': self._is_bula_file(nome)
                    }
                    
                    arquivos_encontrados.append(arquivo_info)
                    self.logger.info(f"🎯 Match NOME COMPLETO encontrado: {nome}")
            
            # 5. Ordenar por relevância (match exato primeiro)
            arquivos_encontrados.sort(key=lambda x: x['score'], reverse=True)
            
            if arquivos_encontrados:
                self.logger.info(f"[OK] {len(arquivos_encontrados)} arquivo(s) relevante(s) para código '{codigo_interno}'")
            else:
                self.logger.warning(f"[AVISO] Nenhum arquivo encontrado para código '{codigo_interno}'")
                
                # Debug: Mostrar amostra dos arquivos analisados (primeiros 5 apenas)
                if arquivos_relevantes:
                    self.logger.debug(f"📋 Amostra de arquivos analisados (primeiros 5):")
                    for item in arquivos_relevantes[:5]:
                        nome_debug = item.get("name", "")
                        codigo_debug = nome_debug.rsplit('.', 1)[0].split('_')[0].split(' ')[0].split('-')[0].strip()
                        self.logger.debug(f"   - {nome_debug} → código extraído: '{codigo_debug}'")
            
            return arquivos_encontrados
            
        except Exception as e:
            self.logger.error(f"💥 Erro geral na busca: {str(e)}")
            return []
    
    def _detectar_tipo_arquivo(self, nome: str) -> str:
        """
        Detecta o tipo de arquivo pela extensão
        
        Args:
            nome: Nome do arquivo
            
        Returns:
            'word', 'image' ou 'pdf'
        """
        nome_lower = nome.lower()
        if nome_lower.endswith('.docx'):
            return 'word'
        elif nome_lower.endswith(('.jpg', '.jpeg', '.png')):
            return 'image'
        elif nome_lower.endswith('.pdf'):
            return 'pdf'
        return 'unknown'
    
    def _is_bula_file(self, nome: str) -> bool:
        """
        Detecta se o arquivo é uma bula ou manual de instalação
        
        Padrões suportados:
        - 10316_BULA_01.jpg
        - 10316_BULA_02.jpg
        - 10836_MANUAL_01.pdf
        - CODIGO_BULA.png
        - CODIGO_MANUAL.docx
        
        Args:
            nome: Nome do arquivo
            
        Returns:
            True se for bula/manual, False caso contrário
        """
        nome_lower = nome.lower()
        
        # Padrões que indicam bula/manual
        keywords_bula = ['bula', 'manual', 'instruction', 'guide', 'info']
        
        # Verificar se contém alguma keyword de bula
        return any(keyword in nome_lower for keyword in keywords_bula)
    
    def download_file_content(self, download_url: str) -> bytes:
        """Baixa conteúdo de um arquivo do SharePoint"""
        try:
            response = requests.get(download_url)
            if response.status_code == 200:
                return response.content
            else:
                raise Exception(f"Erro ao baixar arquivo: {response.status_code}")
                
        except Exception as e:
            raise Exception(f"Erro no download: {str(e)}")
    
    def get_word_images(self, word_file_url: str) -> List[bytes]:
        """
        Extrai imagens de um arquivo Word do SharePoint
        
        Returns:
            Lista de imagens em bytes
        """
        try:
            from docx import Document
            import zipfile
            
            # Baixar arquivo Word
            content = self.download_file_content(word_file_url)
            
            # Processar como documento Word
            word_stream = io.BytesIO(content)
            doc = Document(word_stream)
            
            # Extrair imagens
            imagens = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    img_data = rel.target_part.blob
                    imagens.append(img_data)
                    
                    # Máximo 2 imagens por produto
                    if len(imagens) >= 2:
                        break
            
            return imagens
            
        except Exception as e:
            print(f"[ERRO] Erro ao extrair imagens do Word: {str(e)}")
            return []
    
    def get_word_text(self, word_file_url: str) -> str:
        """
        Extrai texto de um arquivo Word do SharePoint
        
        Returns:
            Texto das especificações técnicas
        """
        try:
            from docx import Document
            
            # Baixar arquivo Word
            content = self.download_file_content(word_file_url)
            
            # Processar documento
            word_stream = io.BytesIO(content)
            doc = Document(word_stream)
            
            # Extrair texto, filtrando linhas vazias e decorativas
            texto_paragrafos = []
            for paragrafo in doc.paragraphs:
                texto = paragrafo.text.strip()
                if texto and not all(c in '_-—–=' for c in texto):
                    texto_paragrafos.append(texto)
            
            # Juntar texto e limitar a 700 caracteres
            texto_completo = '\n'.join(texto_paragrafos)
            if len(texto_completo) > 700:
                texto_completo = texto_completo[:700] + "..."
            
            return texto_completo
            
        except Exception as e:
            print(f"[ERRO] Erro ao extrair texto do Word: {str(e)}")
            return ""


# Exemplo de uso
if __name__ == "__main__":
    client = SharePointClient()
    
    # Testar conexão
    print("[TESTE] Testando conexão SharePoint...")
    print(f"[CAMINHO] {client.PASTA_FICHAS_UNICA}")
    
    # Testar busca por código interno
    print(f"\n[TESTE] Testando busca por código interno...")
    arquivos = client.search_files_by_code("10289")
    print(f"[OK] Arquivos encontrados: {len(arquivos)}")
    for arq in arquivos[:3]:  # Mostrar apenas 3 primeiros
        print(f"   - {arq['name']} (score: {arq.get('score', 0)})")